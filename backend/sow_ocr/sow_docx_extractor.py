"""
sow_docx_extractor.py — Extract text from Word (.docx) files
=============================================================
Uses python-docx to extract paragraphs and tables from Word documents.
Returns plain text in the same format as the PDF extractor so the
rest of the pipeline works identically for both file types.
"""


def extract_text_from_docx(path: str) -> str:
    """
    Extract all text from a .docx file.
    Handles paragraphs and tables.
    Returns a single string — same format as extract_text_from_pdf().
    """
    try:
        from docx import Document
    except ImportError:
        print("  [ERROR] python-docx not installed. Run: pip install python-docx")
        return ""

    try:
        doc = Document(path)
        lines = []

        # ── Extract paragraphs ────────────────────────────────
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # ── Extract tables as "Label   Value" lines ───────────
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if len(cells) == 2:
                    # Key-value row — same format as PDF table extractor
                    lines.append(f"{cells[0]} {cells[1]}")
                elif cells:
                    lines.append("  ".join(cells))

        text = "\n".join(lines)
        print(f"  {len(doc.paragraphs)} paragraphs, {len(text):,} chars extracted")
        return text

    except Exception as e:
        print(f"  [ERROR] Failed to read DOCX: {e}")
        return ""