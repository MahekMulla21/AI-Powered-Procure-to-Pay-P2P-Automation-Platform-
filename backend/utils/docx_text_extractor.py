import docx
import logging

logger = logging.getLogger("docx_text_extractor")

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a .docx file, including paragraphs and tables.
    """
    try:
        doc = docx.Document(file_path)
        full_text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        return ""
